import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement, ns
except ImportError:
    print("python-docx not installed.")
    sys.exit(1)

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')
    
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def main():
    doc = Document()
    
    # 移除底色、設定字體為微軟正黑體
    style = doc.styles['Normal']
    if 'Microsoft JhengHei' not in style.font.name:
        style.font.name = '微軟正黑體'
    style.font.size = Pt(12)
    
    # 加入頁首頁尾與頁碼
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    add_page_number(footer_para)
    
    files = [
        r"c:\資料\AI\waldorf-teacher-os\workspaces\Working_Member\Teacher_林雅婷\projects\第一天_非洲地形故事與題庫.md",
        r"c:\資料\AI\waldorf-teacher-os\workspaces\Working_Member\Teacher_林雅婷\projects\第二天_非洲氣候農業故事與題庫.md",
        r"c:\資料\AI\waldorf-teacher-os\workspaces\Working_Member\Teacher_林雅婷\projects\第三天_非洲礦業人文故事與題庫.md"
    ]
    
    for idx, filepath in enumerate(files):
        if not os.path.exists(filepath):
            print(f"File not found: {filepath}")
            continue
            
        with open(filepath, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        for line in lines:
            text = line.strip()
            if not text or text == '---':
                continue
                
            if text.startswith('### '):
                p = doc.add_heading(text[4:], level=3)
            elif text.startswith('## '):
                p = doc.add_heading(text[3:], level=2)
            elif text.startswith('# '):
                p = doc.add_heading(text[2:], level=1)
            else:
                p = doc.add_paragraph()
                # 處理粗體標籤 **bold**
                parts = re.split(r'(\*\*.*?\*\*)', line.rstrip('\n'))
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)
        
        # 每一天的結尾加入分頁符號 (最後一天除外)
        if idx < len(files) - 1:
            doc.add_page_break()

    output_path = r"g:\我的雲端硬碟\AI\歐亞非地理備課\非洲三天地理故事與題庫.docx"
    doc.save(output_path)
    print(f"Saved successfully to {output_path}")

if __name__ == '__main__':
    main()
