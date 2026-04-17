#!/usr/bin/env python3
# TeacherOS — 識別 Logo 後處理腳本
# 用途：
#   1. 精準偵測圓形邊界，去除白色背景，只保留圓形 Logo
#   2. 壓縮至適合標題行使用的尺寸
#   3. 用無框雙欄表格排版：左欄標題文字、右欄 Logo 垂直置中
#
# 呼叫方式（由 build.sh 自動執行）：
#   python3 setup/add-logo.py <docx路徑>

import sys
from pathlib import Path
from collections import deque

try:
    from docx import Document
    from docx.shared import Cm, Twips
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # 擴展 namespace 支援
    from docx.oxml.ns import nsmap
    if 'wp' not in nsmap:
        nsmap['wp'] = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
        nsmap['a'] = 'http://schemas.openxmlformats.org/drawingml/2006/main'
except ImportError:
    print("錯誤：缺少 python-docx 套件。請執行：pip3 install python-docx")
    sys.exit(1)

try:
    from PIL import Image, ImageDraw
    import numpy as np
    import io
except ImportError:
    print("錯誤：缺少 Pillow 或 numpy 套件。請執行：pip3 install Pillow numpy")
    sys.exit(1)

REPO_ROOT   = Path(__file__).parent.parent
LOGO_SRC    = REPO_ROOT / "setup" / "assets" / "logo.png"
LOGO_READY  = REPO_ROOT / "setup" / "assets" / "logo-ready.png"

LOGO_SIZE_PX    = 300   # 最終圓形 Logo 解析度（px）
LOGO_DEFAULT_PT = 22    # 字體大小預設值（pt）


def flood_fill_background(arr_gray, tolerance=30):
    """
    從四個角落進行 BFS flood fill，偵測白色背景區域。
    回傳 bool mask：True = 背景（要去除），False = 圓形內容（要保留）
    """
    h, w = arr_gray.shape
    visited = np.zeros((h, w), dtype=bool)
    is_background = np.zeros((h, w), dtype=bool)
    queue = deque()

    seeds = [(0, 0), (0, w-1), (h-1, 0), (h-1, w-1)]
    for r, c in seeds:
        if arr_gray[r, c] > (255 - tolerance):
            queue.append((r, c))
            visited[r, c] = True

    while queue:
        r, c = queue.popleft()
        is_background[r, c] = True
        for dr, dc in [(-1,0),(1,0),(0,-1),(0,1)]:
            nr, nc = r+dr, c+dc
            if 0 <= nr < h and 0 <= nc < w and not visited[nr, nc]:
                if arr_gray[nr, nc] > (255 - tolerance):
                    visited[nr, nc] = True
                    queue.append((nr, nc))

    return is_background


def prepare_logo():
    """
    精準去背：
    1. BFS 從角落填充白色背景 → 設為透明
    2. 裁切成正方形並縮放
    3. 套用精確圓形遮罩（確保邊緣乾淨）
    4. 輸出壓縮後的 logo-ready.png
    """
    img = Image.open(LOGO_SRC).convert("RGBA")
    arr = np.array(img)

    gray = np.array(img.convert("L"))
    bg_mask = flood_fill_background(gray, tolerance=25)
    arr[bg_mask, 3] = 0
    img = Image.fromarray(arr)

    alpha = np.array(img.split()[3])
    rows = np.any(alpha > 0, axis=1)
    cols = np.any(alpha > 0, axis=0)
    rmin, rmax = np.where(rows)[0][[0, -1]]
    cmin, cmax = np.where(cols)[0][[0, -1]]
    img = img.crop((cmin, rmin, cmax+1, rmax+1))

    side = max(img.width, img.height)
    square = Image.new("RGBA", (side, side), (0, 0, 0, 0))
    ox = (side - img.width)  // 2
    oy = (side - img.height) // 2
    square.paste(img, (ox, oy), mask=img)
    square = square.resize((LOGO_SIZE_PX, LOGO_SIZE_PX), Image.LANCZOS)

    inset = 3
    circle_mask = Image.new("L", (LOGO_SIZE_PX, LOGO_SIZE_PX), 0)
    draw = ImageDraw.Draw(circle_mask)
    draw.ellipse(
        (inset, inset, LOGO_SIZE_PX - inset - 1, LOGO_SIZE_PX - inset - 1),
        fill=255
    )
    r, g, b, a = square.split()
    final_alpha = Image.fromarray(
        np.minimum(np.array(a), np.array(circle_mask)).astype("uint8")
    )
    result = Image.merge("RGBA", (r, g, b, final_alpha))
    result.save(LOGO_READY, "PNG", optimize=True)

    size_kb = LOGO_READY.stat().st_size // 1024
    arr_check = np.array(result)
    corner_alpha_max = max(
        arr_check[:5,  :5,  3].max(),
        arr_check[:5,  -5:, 3].max(),
        arr_check[-5:, :5,  3].max(),
        arr_check[-5:, -5:, 3].max(),
    )
    print(f"Logo 處理完成：{size_kb}KB（圓形精準去背，四角 alpha={corner_alpha_max}）")
    return LOGO_READY


def logo_to_bytes(logo_path):
    buf = io.BytesIO()
    with Image.open(logo_path) as img:
        img.save(buf, format="PNG")
    buf.seek(0)
    return buf



def detect_font_pt(para):
    """從段落 run 與樣式鏈取得字體大小（pt）。"""
    pt = None
    for run in para.runs:
        if run.font.size:
            pt = run.font.size.pt
            break
    if pt is None:
        style = para.style
        while style and pt is None:
            if style.font.size:
                pt = style.font.size.pt
            style = style.base_style
    return pt or LOGO_DEFAULT_PT



def clear_header(doc):
    """清除所有 section 的頁首內容"""
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        for para in header.paragraphs:
            para.clear()


def add_watermark_background(doc, bg_path_str):
    """
    將背景圖片作為全頁浮水印插入每頁頁首。
    使用 VML 格式（Word 原生浮水印機制），Google Docs 相容性較好。
    """
    from docx.shared import Inches
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    import hashlib

    bg_path = Path(bg_path_str)
    if not bg_path.exists():
        print(f"警告：找不到背景圖片 {bg_path}，略過浮水印邏輯。")
        return

    # A4 尺寸（pt）
    page_w_pt = 595
    page_h_pt = 842

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False

        # 將背景圖片加入 header 的 relationship
        rId = header.part.relate_to(
            doc.part.package.get_or_add_image_part(str(bg_path)),
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'
        )

        # 清除既有內容
        for para in header.paragraphs:
            para.clear()

        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

        # 建立 VML 浮水印（Word 原生格式）
        run = para.add_run()
        pict = OxmlElement('w:pict')

        # VML shape：全頁置中、文字後方
        shape_xml = (
            f'<v:shape xmlns:v="urn:schemas-microsoft-com:vml" '
            f'xmlns:o="urn:schemas-microsoft-com:office:office" '
            f'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
            f'id="WordPictureWatermark" o:spid="_x0000_s2049" '
            f'type="#_x0000_t75" '
            f'style="position:absolute;margin-left:0;margin-top:0;'
            f'width:{page_w_pt}pt;height:{page_h_pt}pt;z-index:-251658752;'
            f'mso-position-horizontal:center;mso-position-horizontal-relative:margin;'
            f'mso-position-vertical:center;mso-position-vertical-relative:margin" '
            f'o:allowincell="f">'
            f'<v:imagedata r:id="{rId}" o:title="waldorf-bg"/>'
            f'</v:shape>'
        )

        from lxml import etree
        shape_el = etree.fromstring(shape_xml)
        pict.append(shape_el)
        run._r.append(pict)


def add_logo_inline_layout(docx_path: str, logo_path: Path):
    """
    以 Inline 方式插入 Logo，避免表格切割文字。
    為了垂直置中 Logo，我們改為「提昇標題文字」的 baseline (w:position)，
    讓 Logo 維持在其原本基準網格。
    """
    doc = Document(docx_path)
    clear_header(doc)

    title_para = None
    for para in doc.paragraphs:
        if para.text.strip():
            title_para = para
            break

    if title_para is None:
        print("警告：找不到標題段落，略過 Logo 嵌入。")
        doc.save(docx_path)
        return

    font_pt   = detect_font_pt(title_para)
    logo_h_cm = max(font_pt * 0.06, 1.2)
    logo_h_pt = logo_h_cm * 28.346
    
    # 預估標題字體的視覺高度 (Cap height) 約為 70% pt
    cap_height = font_pt * 0.7
    
    # 計算需要提升文字多少點，才能讓文字中心對齊 Logo 中心
    diff_pt = max(0, (logo_h_pt / 2.0) - (cap_height / 2.0))
    pos_val = int(diff_pt * 2)  # half-points for OpenXML

    # 提昇所有既有文字 run
    for run in title_para.runs:
        rPr = run._r.get_or_add_rPr()
        pos_el = OxmlElement('w:position')
        pos_el.set(qn('w:val'), str(pos_val))
        rPr.append(pos_el)

    # 加兩個空白格，一樣要提升
    space_run = title_para.add_run("  ")
    rPr_space = space_run._r.get_or_add_rPr()
    pos_space = OxmlElement('w:position')
    pos_space.set(qn('w:val'), str(pos_val))
    rPr_space.append(pos_space)

    # 插入 Logo (不套用偏移，讓他直立於 Paragraph 基準線)
    logo_run = title_para.add_run()
    logo_run.add_picture(logo_to_bytes(logo_path), height=Cm(logo_h_cm))

    doc.save(docx_path)
    print(
        f"Logo 已嵌入標題末端（高度 {logo_h_cm:.2f}cm，Inline 模式，文字基線偏移 {diff_pt:.1f}pt）："
        f"{Path(docx_path).name}"
    )


def main():
    if len(sys.argv) < 2:
        print("用法：python3 setup/add-logo.py <docx路徑>")
        sys.exit(1)

    docx_path = sys.argv[1]

    if not Path(docx_path).exists():
        print(f"錯誤：找不到檔案 {docx_path}")
        sys.exit(1)

    if not LOGO_SRC.exists():
        print(f"警告：找不到 Logo 原始檔（{LOGO_SRC}），略過。")
        return

    logo_ready = prepare_logo()
    add_logo_inline_layout(docx_path, logo_ready)
    
    # 加入華德福水彩背景
    bg_jpg_path = REPO_ROOT / "setup" / "assets" / "waldorf-bg.jpg"
    doc = Document(docx_path)
    add_watermark_background(doc, str(bg_jpg_path))
    doc.save(docx_path)
    print(f"已加入華德福淡彩水彩浮水印背景：{Path(docx_path).name}")


if __name__ == "__main__":
    main()
